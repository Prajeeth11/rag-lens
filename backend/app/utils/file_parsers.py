from dataclasses import dataclass, field
from pathlib import Path


@dataclass
class ParsedDocument:
    text: str
    pages: list[str] = field(default_factory=list)
    metadata: dict = field(default_factory=dict)


def parse_file(path: str | Path) -> ParsedDocument:
    path = Path(path)
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _parse_pdf(path)
    if suffix == ".docx":
        return _parse_docx(path)
    if suffix in (".txt", ".md", ".markdown"):
        text = path.read_text(encoding="utf-8", errors="replace")
        return ParsedDocument(text=text, pages=[text], metadata={"format": suffix.lstrip(".")})
    raise ValueError(f"Unsupported file type: {suffix}")


def _parse_pdf(path: Path) -> ParsedDocument:
    import fitz  # PyMuPDF

    pages: list[str] = []
    with fitz.open(path) as doc:
        for page in doc:
            pages.append(page.get_text())
        meta = {"format": "pdf", "num_pages": len(doc), "title": doc.metadata.get("title") or path.stem}
    return ParsedDocument(text="\n\n".join(pages), pages=pages, metadata=meta)


def _parse_docx(path: Path) -> ParsedDocument:
    import docx

    doc = docx.Document(str(path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    text = "\n\n".join(paragraphs)
    return ParsedDocument(text=text, pages=[text], metadata={"format": "docx", "num_paragraphs": len(paragraphs)})
